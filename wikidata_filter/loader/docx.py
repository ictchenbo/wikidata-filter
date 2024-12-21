"""
基于python-docx读取docx文件 输出文本段落及表格

{"meta": {}, "type": "paragraph", "content": "paragraph content"}
{"meta": {}, "type": "table", "content": [[]] }

@dependencies pdfminer
"""
import os
from datetime import datetime
from typing import Iterable, Any
from wikidata_filter.loader.file import BinaryFile

try:
    from docx import Document
except:
    print("Error! you need to install python-docx first: pip install python-docx")
    raise ImportError("python-docx not installed")


class Docx(BinaryFile):
    """基于python-docx读取docx文件"""
    def __init__(self, input_file: str, max_pages: int = 0, **kwargs):
        super().__init__(input_file, **kwargs)
        self.max_pages = max_pages

    def iter(self) -> Iterable[Any]:
        document = Document(self.instream)
        core = document.core_properties
        props = {}

        def set_attr(k, v=None):
            if v:
                if isinstance(v, datetime):
                    v = v.strftime('%Y-%m-%d %H:%M:%S')
                props[k] = v

        set_attr('title', core.title)
        set_attr('subject', core.subject)
        set_attr('title', core.title)
        set_attr('subject', core.subject)
        set_attr('author', core.author)
        set_attr('keywords', core.keywords)
        set_attr('comments', core.comments)
        set_attr('created', core.created)
        set_attr('modified', core.modified)
        set_attr('language', core.language)
        set_attr('identifier', core.identifier)
        set_attr('version', core.version)
        set_attr('category', core.category)

        for paragraph in document.paragraphs:
            yield {
                "meta": props,
                "type": "paragraph",
                "style": paragraph.style.name,
                "content": paragraph.text
            }
        for table in document.tables:
            rows = [
                [cell.text.strip() for cell in row.cells]
                for row in table.rows
            ]
            yield {
                "meta": props,
                "type": "table",
                "content": rows
            }


class Doc(Docx):
    """基于libreoffice6将doc转换为docx 进而基于docx进行解析"""
    def __init__(self, input_file: str, max_pages: int = 0, **kwargs):
        super().__init__(input_file, max_pages=max_pages, auto_open=False, **kwargs)
        out_path = os.path.dirname(self.filename)
        try:
            self.doc2docx(self.filename, out_path)
        except:
            print('doc2docx failed! make sure you have install libreoffice6.4')
        self.filename = self.filename + "x"
        self.instream = open(self.filename, 'rb')

    def doc2docx(self, doc: str, output_dir: str):
        os.system(f'libreoffice6.4 --headless --convert-to docx "{doc}" --outdir "{output_dir}"')
